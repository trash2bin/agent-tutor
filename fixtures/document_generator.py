from __future__ import annotations

import base64
import json
import os
import re
import textwrap
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Protocol, Sequence

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt

from db.database import Database, PROJECT_ROOT
from db.models import Discipline, Material
from tools.rag import RagTools


DEFAULT_DOCGEN_MODEL = "qwen2.5:0.5b"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "generated_materials"
DOCX_FONT_NAME = "Times New Roman"
DOCX_BODY_SIZE_PT = 14


@dataclass(frozen=True)
class MaterialSpec:
    material_type: str
    extension: str
    title_prefix: str

    @property
    def suffix(self) -> str:
        return f".{self.extension}"


@dataclass(frozen=True)
class GeneratedDocument:
    title: str
    material_type: str
    path: Path
    text: str


@dataclass(frozen=True)
class OllamaGenerateConfig:
    model: str
    url: str
    timeout_seconds: float
    temperature: float
    num_predict: int
    min_response_chars: int

    @classmethod
    def from_env(cls) -> OllamaGenerateConfig:
        return cls(
            model=os.environ.get("DOCGEN_MODEL", DEFAULT_DOCGEN_MODEL),
            url=os.environ.get("DOCGEN_OLLAMA_URL", _default_ollama_generate_url()),
            timeout_seconds=_env_float("DOCGEN_TIMEOUT", 90.0),
            temperature=_env_float("DOCGEN_TEMPERATURE", 0.35),
            num_predict=_env_int("DOCGEN_NUM_PREDICT", 4500),
            min_response_chars=_env_int("DOCGEN_MIN_RESPONSE_CHARS", 300),
        )


class TextGenerationClient(Protocol):
    def generate(
        self,
        prompt: str,
        images: Sequence[str | Path] | None = None,
    ) -> str:
        """Generate text from a prompt and optional image inputs."""


class OllamaGenerateClient:
    """Small client for Ollama's /api/generate endpoint.

    Ollama accepts base64-encoded images for multimodal models. The document
    generator does not need images today, but keeping the hook here makes that a
    client concern instead of leaking HTTP payload details into generation code.
    """

    def __init__(self, config: OllamaGenerateConfig | None = None) -> None:
        self.config = config or OllamaGenerateConfig.from_env()

    def generate(
        self,
        prompt: str,
        images: Sequence[str | Path] | None = None,
    ) -> str:
        payload = self._payload(prompt, images)
        request = urllib.request.Request(
            self.config.url,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )

        try:
            with urllib.request.urlopen(
                request,
                timeout=self.config.timeout_seconds,
            ) as response:
                data = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            raise RuntimeError(self._http_error_message(exc)) from exc
        except (OSError, TimeoutError, json.JSONDecodeError) as exc:
            raise RuntimeError(
                "Ollama is unavailable. Start it with `ollama serve` and make "
                f"sure model `{self.config.model}` is installed."
            ) from exc

        if error := data.get("error"):
            raise RuntimeError(f"Ollama returned an error: {error}")

        text = str(data.get("response", "")).strip()
        if len(text) < self.config.min_response_chars:
            raise RuntimeError(
                f"Ollama model `{self.config.model}` returned too little text."
            )
        return text

    def _payload(
        self,
        prompt: str,
        images: Sequence[str | Path] | None,
    ) -> dict[str, Any]:
        payload: dict[str, Any] = {
            "model": self.config.model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": self.config.temperature,
                "num_predict": self.config.num_predict,
            },
        }
        if images:
            payload["images"] = [_image_to_base64(image) for image in images]
        return payload

    def _http_error_message(self, exc: urllib.error.HTTPError) -> str:
        details = exc.read().decode("utf-8", errors="replace").strip()
        if details:
            return f"Ollama request failed with HTTP {exc.code}: {details}"
        return f"Ollama request failed with HTTP {exc.code}."


class DocumentTextGenerator:
    """Builds educational prompts and delegates text generation to a client."""

    def __init__(self, client: TextGenerationClient | None = None) -> None:
        self.client = client or OllamaGenerateClient()

    def generate(self, discipline: Discipline, material_type: str) -> str:
        return self.client.generate(self.build_prompt(discipline, material_type))

    @staticmethod
    def build_prompt(discipline: Discipline, material_type: str) -> str:
        lines = [
            "Сгенерируй учебный материал для университета на русском языке.",
            f"Дисциплина: {discipline.name}",
            f"Описание дисциплины: {discipline.description}",
            f"Тип материала: {material_type}",
            "Требования:",
            "- объём 1200-1800 слов;",
            "- структурированный учебный текст с разделами, подпунктами и примерами;",
            "- используй Markdown только для заголовков (#, ##, ###), списков "
            "и **жирного текста**;",
            "- добавь практический пример, контрольные вопросы и краткие выводы;",
            "- без выдуманных ссылок и без обращения к читателю от лица ассистента.",
        ]
        return "\n".join(lines)


MATERIAL_SPECS: tuple[MaterialSpec, ...] = (
    MaterialSpec("Лекция", "pdf", "Лекция"),
    MaterialSpec("Методичка", "docx", "Методичка"),
    MaterialSpec(
        "Лабораторная работа",
        "docx",
        "Лабораторная работа",
    ),
)


class MaterialDocumentGenerator:
    def __init__(
        self,
        db: Database,
        rag_tools: RagTools,
        output_dir: str | Path | None = None,
        text_generator: DocumentTextGenerator | None = None,
        material_specs: Sequence[MaterialSpec] = MATERIAL_SPECS,
    ) -> None:
        self.db = db
        self.rag_tools = rag_tools
        self.output_dir = Path(
            output_dir or os.environ.get("DOCGEN_OUTPUT_DIR", DEFAULT_OUTPUT_DIR)
        )
        self.text_generator = text_generator or DocumentTextGenerator()
        self.material_specs = tuple(material_specs)

    def ensure_materials(
        self,
        discipline_id: str,
        force: bool = False,
    ) -> list[Material]:
        discipline = self.db.get_discipline(discipline_id)
        if discipline is None:
            return []

        self._remove_stale_generated_documents(discipline_id, force=force)

        existing = self._valid_generated_materials(
            self.db.get_materials(discipline_id)
        )
        missing_types = self._expected_material_types() - {
            material.type for material in existing
        }
        if existing and not force and not missing_types:
            return existing

        generated = self.generate_documents(
            discipline,
            material_types=None if force else missing_types,
        )
        for document in generated:
            self._index_generated_document(discipline.id, document)

        return self.db.get_materials(discipline_id)

    def generate_documents(
        self,
        discipline: Discipline,
        material_types: set[str] | None = None,
    ) -> list[GeneratedDocument]:
        discipline_dir = self.output_dir / _slugify(discipline.name)
        discipline_dir.mkdir(parents=True, exist_ok=True)

        generated: list[GeneratedDocument] = []
        for spec in self.material_specs:
            if material_types is not None and spec.material_type not in material_types:
                continue
            generated.append(self._generate_document(discipline, spec, discipline_dir))
        return generated

    def _generate_document(
        self,
        discipline: Discipline,
        spec: MaterialSpec,
        discipline_dir: Path,
    ) -> GeneratedDocument:
        title = f"{spec.title_prefix}: {discipline.name}"
        text = self.text_generator.generate(discipline, spec.material_type)
        path = discipline_dir / self._file_name(spec, discipline)

        _write_document(path, title, text)
        return GeneratedDocument(
            title=title,
            material_type=spec.material_type,
            path=path,
            text=text,
        )

    def _file_name(self, spec: MaterialSpec, discipline: Discipline) -> str:
        prefix = _slugify(spec.title_prefix)
        discipline_name = _slugify(discipline.name)
        return f"{prefix}_{discipline_name}{spec.suffix}"

    def _index_generated_document(
        self,
        discipline_id: str,
        document: GeneratedDocument,
    ) -> None:
        try:
            chunks = self.rag_tools.pipeline.chunker.chunk_pages(
                [{"page": None, "text": document.text}]
            )
            self.rag_tools.pipeline._save_document(
                source_path=document.path.resolve(),
                chunks=chunks,
                discipline_id=discipline_id,
                title=document.title,
            )
        except Exception:
            self.db.save_generated_document_record(
                path=str(document.path),
                discipline_id=discipline_id,
                title=document.title,
                text=document.text,
            )

    def _remove_stale_generated_documents(
        self,
        discipline_id: str,
        force: bool,
    ) -> None:
        if force:
            self._delete_generated_documents_where(
                discipline_id=discipline_id,
                missing_only=False,
            )
            return

        self._delete_generated_documents_where(
            discipline_id=discipline_id,
            missing_only=True,
        )
        self._delete_outdated_generated_documents(discipline_id)

    def _delete_outdated_generated_documents(self, discipline_id: str) -> None:
        expected_extensions = self._expected_extensions()
        rows_to_delete = []
        for material in self.db.get_materials(discipline_id):
            source_path = Path(material.source_path)
            if not self._is_generated_path(source_path):
                continue

            expected_extension = expected_extensions.get(material.type)
            if expected_extension and source_path.suffix.lower() != expected_extension:
                rows_to_delete.append(
                    {"id": material.id, "source_path": material.source_path}
                )

        self._delete_document_rows(rows_to_delete)

    def _delete_generated_documents_where(
        self,
        discipline_id: str,
        missing_only: bool,
    ) -> None:
        cursor = self.db.conn.cursor()
        rows = cursor.execute(
            """
            SELECT id, source_path FROM documents
            WHERE discipline_id = ? AND source_path LIKE ?
            """,
            (discipline_id, f"%{self.output_dir.name}%"),
        ).fetchall()

        rows_to_delete = []
        for row in rows:
            source_path = Path(row["source_path"])
            if not self._is_generated_path(source_path):
                continue
            if missing_only and source_path.exists():
                continue
            rows_to_delete.append(row)

        self._delete_document_rows(rows_to_delete)

    def _delete_document_rows(self, rows: Iterable[Any]) -> None:
        cursor = self.db.conn.cursor()
        deleted_any = False
        for row in rows:
            self._delete_generated_document_row(cursor, row)
            deleted_any = True

        if deleted_any:
            self.db.conn.commit()
            self._cleanup_empty_output_dirs()

    def _delete_generated_document_row(self, cursor: Any, row: Any) -> None:
        document_id = row["id"]
        source_path = Path(row["source_path"])

        try:
            self.rag_tools._delete_document_vectors(document_id)
        except Exception:
            pass

        cursor.execute(
            "DELETE FROM document_chunks WHERE document_id = ?",
            (document_id,),
        )
        cursor.execute("DELETE FROM documents WHERE id = ?", (document_id,))
        if source_path.exists():
            try:
                source_path.unlink()
            except OSError:
                pass

    def _cleanup_empty_output_dirs(self) -> None:
        if not self.output_dir.exists():
            return

        paths = sorted(
            self.output_dir.rglob("*"),
            key=lambda item: len(item.parts),
            reverse=True,
        )
        for path in paths:
            if not path.is_dir():
                continue
            try:
                path.rmdir()
            except OSError:
                pass

        try:
            self.output_dir.rmdir()
        except OSError:
            pass

    def _valid_generated_materials(self, materials: list[Material]) -> list[Material]:
        expected_extensions = self._expected_extensions()
        valid = []
        for material in materials:
            source_path = Path(material.source_path)
            expected_extension = expected_extensions.get(material.type)
            if (
                expected_extension
                and self._is_generated_path(source_path)
                and source_path.suffix.lower() != expected_extension
            ):
                continue
            valid.append(material)
        return valid

    def _expected_extensions(self) -> dict[str, str]:
        return {spec.material_type: spec.suffix for spec in self.material_specs}

    def _expected_material_types(self) -> set[str]:
        return {spec.material_type for spec in self.material_specs}

    def _is_generated_path(self, source_path: Path) -> bool:
        if self.output_dir.name not in source_path.parts:
            return False

        try:
            source_path.resolve().relative_to(self.output_dir.resolve())
            return True
        except (OSError, ValueError):
            return False


def _default_ollama_generate_url() -> str:
    host = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434").rstrip("/")
    if not host.startswith(("http://", "https://")):
        host = f"http://{host}"
    return f"{host}/api/generate"


def _env_int(name: str, default: int) -> int:
    value = os.environ.get(name)
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    value = os.environ.get(name)
    if value is None:
        return default
    try:
        return float(value)
    except ValueError:
        return default


def _image_to_base64(image: str | Path) -> str:
    image_path = Path(image)
    try:
        exists = image_path.exists()
    except OSError:
        exists = False
    if exists:
        return base64.b64encode(image_path.read_bytes()).decode("ascii")
    return str(image)


def _slugify(value: str) -> str:
    value = value.lower().replace("ё", "е")
    value = "".join(TRANSLIT.get(char, char) for char in value)
    value = re.sub(r"[^a-z0-9_]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "material"


TRANSLIT = {
    "а": "a",
    "б": "b",
    "в": "v",
    "г": "g",
    "д": "d",
    "е": "e",
    "ж": "zh",
    "з": "z",
    "и": "i",
    "й": "y",
    "к": "k",
    "л": "l",
    "м": "m",
    "н": "n",
    "о": "o",
    "п": "p",
    "р": "r",
    "с": "s",
    "т": "t",
    "у": "u",
    "ф": "f",
    "х": "h",
    "ц": "ts",
    "ч": "ch",
    "ш": "sh",
    "щ": "sch",
    "ъ": "",
    "ы": "y",
    "ь": "",
    "э": "e",
    "ю": "yu",
    "я": "ya",
    " ": "_",
}


def _write_document(path: Path, title: str, text: str) -> None:
    if path.suffix == ".docx":
        _write_docx(path, title, text)
        return
    if path.suffix == ".pdf":
        _write_pdf(path, title, text)
        return
    raise ValueError(f"Unsupported generated document extension: {path.suffix}")


def _write_docx(path: Path, title: str, text: str) -> None:
    document = DocxDocument()
    _configure_docx_styles(document)

    title_paragraph = document.add_heading("", 0)
    _add_markdown_runs(title_paragraph, title)

    for block in _markdown_blocks(text):
        _add_docx_block(document, block)

    document.save(path)


def _add_docx_block(document: Any, block: str) -> None:
    heading_match = re.match(r"^(#{1,4})\s+(.+)$", block)
    if heading_match:
        level = min(len(heading_match.group(1)), 3)
        paragraph = document.add_heading("", level=level)
        _add_markdown_runs(paragraph, heading_match.group(2))
        return

    if re.match(r"^[-*]\s+", block):
        paragraph = document.add_paragraph(style="List Bullet")
        _add_markdown_runs(paragraph, re.sub(r"^[-*]\s+", "", block))
        return

    if re.match(r"^\d+[.)]\s+", block):
        paragraph = document.add_paragraph(style="List Number")
        _add_markdown_runs(paragraph, re.sub(r"^\d+[.)]\s+", "", block))
        return

    paragraph = document.add_paragraph()
    _add_markdown_runs(paragraph, block)


def _configure_docx_styles(document: Any) -> None:
    for style_name in ["Normal", "List Bullet", "List Number"]:
        _configure_docx_font(document.styles[style_name], DOCX_BODY_SIZE_PT)

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 16, True),
        ("Heading 2", 15, True),
        ("Heading 3", 14, True),
    ]:
        _configure_docx_font(document.styles[style_name], size, bold=bold)


def _configure_docx_font(style: Any, size: int, bold: bool = False) -> None:
    style.font.name = DOCX_FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)


def _add_markdown_runs(paragraph: Any, text: str) -> None:
    for part, is_bold in _split_bold(text):
        run = paragraph.add_run(_strip_inline_markdown(part))
        run.bold = is_bold
        run.font.name = DOCX_FONT_NAME
        run.font.size = Pt(DOCX_BODY_SIZE_PT)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)


def _write_pdf(path: Path, title: str, text: str) -> None:
    wrapped_lines = [title, ""]
    for paragraph in _markdown_blocks(text):
        clean_paragraph = _strip_markdown(paragraph)
        wrapped_lines.extend(textwrap.wrap(clean_paragraph, width=74) or [""])
        wrapped_lines.append("")

    pages = list(_paginate(wrapped_lines, lines_per_page=36))
    objects = _build_pdf_objects(pages)
    _write_pdf_objects(path, objects)


def _build_pdf_objects(pages: list[list[str]]) -> list[bytes]:
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    page_refs = " ".join(f"{3 + index * 2} 0 R" for index in range(len(pages)))
    objects.append(
        f"<< /Type /Pages /Kids [{page_refs}] /Count {len(pages)} >>".encode()
    )

    font_object_id = 3 + len(pages) * 2
    for index, page_lines in enumerate(pages):
        page_id = 3 + index * 2
        content_id = page_id + 1
        objects.append(_pdf_page_object(page_id, content_id, font_object_id))
        objects.append(_pdf_stream_object(_pdf_text_stream(page_lines)))

    objects.extend(_pdf_font_objects(font_object_id))
    return objects


def _pdf_page_object(page_id: int, content_id: int, font_object_id: int) -> bytes:
    return (
        f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        f"/Resources << /Font << /F1 {font_object_id} 0 R >> >> "
        f"/Contents {content_id} 0 R >>"
    ).encode()


def _pdf_stream_object(stream: bytes) -> bytes:
    return (
        b"<< /Length "
        + str(len(stream)).encode()
        + b" >>\nstream\n"
        + stream
        + b"\nendstream"
    )


def _pdf_font_objects(font_object_id: int) -> list[bytes]:
    cmap = _to_unicode_cmap()
    return [
        (
            b"<< /Type /Font /Subtype /Type0 /BaseFont /Times-Roman "
            b"/Encoding /Identity-H /DescendantFonts ["
            + str(font_object_id + 1).encode()
            + b" 0 R] /ToUnicode "
            + str(font_object_id + 2).encode()
            + b" 0 R >>"
        ),
        (
            b"<< /Type /Font /Subtype /CIDFontType2 /BaseFont /Times-Roman "
            b"/CIDSystemInfo << /Registry (Adobe) /Ordering (Identity) "
            b"/Supplement 0 >> "
            b"/FontDescriptor "
            + str(font_object_id + 3).encode()
            + b" 0 R /W [0 [600]] >>"
        ),
        _pdf_stream_object(cmap),
        (
            b"<< /Type /FontDescriptor /FontName /Times-Roman /Flags 4 "
            b"/FontBBox [-166 -225 1000 931] /ItalicAngle 0 /Ascent 931 "
            b"/Descent -225 /CapHeight 718 /StemV 80 >>"
        ),
    ]


def _markdown_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            blocks.append(" ".join(buffer).strip())
            buffer.clear()

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush_buffer()
            continue
        if re.fullmatch(r"#{1,6}", line):
            flush_buffer()
            continue
        if re.match(r"^(#{1,6}\s*|[-*]\s+|\d+[.)]\s+)", line):
            flush_buffer()
            blocks.append(line)
            continue
        buffer.append(line)

    flush_buffer()
    return blocks


def _split_bold(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    position = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > position:
            parts.append((text[position:match.start()], False))
        parts.append((match.group(1), True))
        position = match.end()
    if position < len(text):
        parts.append((text[position:], False))
    return parts or [(text, False)]


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("***", "").replace("__", "")


def _strip_markdown(text: str) -> str:
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^[-*]\s+", "- ", text)
    text = re.sub(
        r"^\d+[.)]\s+",
        lambda match: match.group(0).replace(")", "."),
        text,
    )
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return _strip_inline_markdown(text)


def _paginate(lines: list[str], lines_per_page: int) -> Iterable[list[str]]:
    for index in range(0, len(lines), lines_per_page):
        yield lines[index : index + lines_per_page]


def _pdf_text_stream(lines: list[str]) -> bytes:
    stream = ["BT", "/F1 14 Tf", "50 790 Td", "18 TL"]
    for line in lines:
        stream.append(f"<{line.encode('utf-16-be').hex()}> Tj")
        stream.append("T*")
    stream.append("ET")
    return "\n".join(stream).encode("ascii")


def _to_unicode_cmap() -> bytes:
    return (
        "/CIDInit /ProcSet findresource begin\n"
        "12 dict begin\n"
        "begincmap\n"
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (UCS) /Supplement 0 >> def\n"
        "/CMapName /Adobe-Identity-UCS def\n"
        "/CMapType 2 def\n"
        "1 begincodespacerange\n"
        "<0000> <FFFF>\n"
        "endcodespacerange\n"
        "1 beginbfrange\n"
        "<0000> <FFFF> <0000>\n"
        "endbfrange\n"
        "endcmap\n"
        "CMapName currentdict /CMap defineresource pop\n"
        "end\n"
        "end"
    ).encode("ascii")


def _write_pdf_objects(path: Path, objects: list[bytes]) -> None:
    content = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")

    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(bytes(content))
